"""
Post-procesa los DOCX rendereados por Quarto para que las tablas se vean
bonitas (Word):
    - cabecera con fondo azul-cyan + texto blanco bold
    - bordes finos grises en todas las celdas
    - texto compacto (10 pt)
    - ancho de tabla = 100% de la página
    - reemplaza emojis de estado 🟢/🟡/🔴 por círculos ● coloreados
      (verde/amarillo/rojo) que sí renderizan en cualquier fuente Word

Uso:
    cd /home/rstudio/work/proyecto-cgsm
    python scripts/embellish_docx_tables.py docs/informe_final.docx
    python scripts/embellish_docx_tables.py docs/informe_anexos.docx

Requiere:  pip install python-docx
"""
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── estilo ─────────────────────────────────────────────────────────────
HEADER_BG       = '2E86AB'   # azul medio (mismo tono que la imagen ref)
HEADER_TEXT     = RGBColor(0xFF, 0xFF, 0xFF)
BODY_TEXT       = RGBColor(0x00, 0x00, 0x00)
BORDER_COLOR    = 'BFBFBF'   # gris claro
BORDER_SIZE     = '4'        # eighths of pt → 4 = 0.5 pt fina
TABLE_FONT_SZ   = Pt(10)
HEADER_FONT_SZ  = Pt(10)

# emojis 🟢🟡🔴 → ● coloreado (Word renderiza ● fiable en cualquier fuente)
STATUS_EMOJI = {
    '🟢': '4CAF50',   # verde
    '🟡': 'FFC107',   # amarillo
    '🔴': 'F44336',   # rojo
}
STATUS_GLYPH = '●'


def _hex_to_rgb(h):
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_cell_borders(cell, color=BORDER_COLOR, size=BORDER_SIZE):
    """Bordes finos grises en los cuatro lados de la celda."""
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn('w:tcBorders')):
        tc_pr.remove(old)
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _strip_table_style(tbl):
    """Quita el estilo heredado del reference doc y reaplica bordes
    finos grises a nivel tabla. Quita posicionamiento flotante."""
    tbl_pr = tbl._tbl.tblPr
    for s in tbl_pr.findall(qn('w:tblStyle')):
        tbl_pr.remove(s)
    for old in tbl_pr.findall(qn('w:tblBorders')):
        tbl_pr.remove(old)
    for pos in tbl_pr.findall(qn('w:tblpPr')):
        tbl_pr.remove(pos)
    for wrap in tbl_pr.findall(qn('w:tblOverlap')):
        tbl_pr.remove(wrap)
    tbl_borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), BORDER_SIZE)
        b.set(qn('w:color'), BORDER_COLOR)
        tbl_borders.append(b)
    tbl_pr.append(tbl_borders)


def _strip_paragraph_borders(doc):
    """Quita cualquier w:pBdr (border-around-paragraph) en body y celdas."""
    def _strip(p):
        p_pr = p._p.find(qn('w:pPr'))
        if p_pr is None:
            return
        for pbdr in p_pr.findall(qn('w:pBdr')):
            p_pr.remove(pbdr)
    for p in doc.paragraphs:
        _strip(p)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _strip(p)


def _compact_layout(doc):
    """Reduce margenes y tamano de fuente Normal para que el DOCX se
    pagine compacto similar al PDF. Sin esto, los defaults de Word
    (11pt + 2.54cm) duplican la cantidad de paginas vs el PDF que
    usa scriptsize global."""
    # Margenes: 1.8 cm top/bot, 2 cm left/right (twips: 1 cm = 567)
    PG_MAR = {
        'top':    str(int(1.8 * 567)),
        'bottom': str(int(1.8 * 567)),
        'left':   str(int(2.0 * 567)),
        'right':  str(int(2.0 * 567)),
        'header': str(int(0.8 * 567)),
        'footer': str(int(0.8 * 567)),
        'gutter': '0',
    }
    for section in doc.sections:
        sectPr = section._sectPr
        # Quitar pgMar previo si existe
        for old in sectPr.findall(qn('w:pgMar')):
            sectPr.remove(old)
        pgMar = OxmlElement('w:pgMar')
        for k, v in PG_MAR.items():
            pgMar.set(qn(f'w:{k}'), v)
        # Insertar pgMar dentro de sectPr
        sectPr.append(pgMar)
    # Reducir fuente del estilo Normal a 10 pt
    try:
        normal = doc.styles['Normal']
        normal.font.size = Pt(10)
    except KeyError:
        pass
    # Tambien comprimir interlineado del Normal (1.15 -> 1.05)
    try:
        normal = doc.styles['Normal']
        pf = normal.paragraph_format
        pf.line_spacing = 1.05
        pf.space_after = Pt(2)
        pf.space_before = Pt(0)
    except (KeyError, AttributeError):
        pass


def _strip_caption_shading(doc):
    """Quita el w:shd (relleno) de cualquier párrafo que esté FUERA de
    una tabla, es decir, de captions de tabla/figura y prosa del body.
    Esto borra la banda cyan que el reference doc aplica a las captions."""
    for p in doc.paragraphs:
        p_pr = p._p.find(qn('w:pPr'))
        if p_pr is None:
            continue
        for shd in p_pr.findall(qn('w:shd')):
            p_pr.remove(shd)


def _strip_figure_borders(doc):
    """Quita el contorno (a:ln) de TODAS las imágenes."""
    A_NS  = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    A_LN  = f'{{{A_NS}}}ln'
    for ln in list(doc.element.iter(A_LN)):
        parent = ln.getparent()
        if parent is not None:
            parent.remove(ln)
    PIC_NS = 'http://schemas.openxmlformats.org/drawingml/2006/picture'
    SP_PR  = f'{{{PIC_NS}}}spPr'
    from lxml import etree
    for sp_pr in doc.element.iter(SP_PR):
        ln_off = etree.SubElement(sp_pr, A_LN)
        etree.SubElement(ln_off, '{%s}noFill' % A_NS)


def _set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn('w:shd')):
        tc_pr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def _clear_cell_shading(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    for shd in tc_pr.findall(qn('w:shd')):
        tc_pr.remove(shd)


def _table_contains_drawing(tbl):
    """True si la tabla contiene una imagen/figura embebida (wrapper de
    figura generado por pandoc). Esas tablas no deben recibir styling
    de header cyan."""
    for el in tbl._tbl.iter():
        tag = el.tag
        if tag.endswith('}drawing') or tag.endswith('}blip') or tag.endswith('}pic'):
            return True
    return False


def _unwrap_figure_tables(doc):
    """Reemplaza cada tabla wrapper de figura por los parrafos que
    contiene su unica celda. Pandoc envuelve cada figura en una tabla
    1x1 para layout; al quitar la tabla y dejar solo los parrafos, el
    DOCX queda sin la 'caja' invisible alrededor de la figura."""
    unwrapped = 0
    for tbl in list(doc.tables):
        if not _table_contains_drawing(tbl):
            continue
        tbl_el = tbl._tbl
        tbl_parent = tbl_el.getparent()
        # Recolectar parrafos de la unica celda (en orden)
        cell_paragraphs = []
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    cell_paragraphs.append(p._p)
        # Insertar parrafos justo antes de la tabla (lxml mueve, no copia)
        tbl_idx = list(tbl_parent).index(tbl_el)
        for i, p_el in enumerate(cell_paragraphs):
            tbl_parent.insert(tbl_idx + i, p_el)
        # Borrar la tabla vacia
        tbl_parent.remove(tbl_el)
        unwrapped += 1
    return unwrapped


def _strip_table_borders_and_shading(tbl):
    """Para tablas wrapper de figuras: borra bordes y shading
    completamente, sin aplicar el styling de header cyan."""
    tbl_pr = tbl._tbl.tblPr
    for s in tbl_pr.findall(qn('w:tblStyle')):
        tbl_pr.remove(s)
    for old in tbl_pr.findall(qn('w:tblBorders')):
        tbl_pr.remove(old)
    for pos in tbl_pr.findall(qn('w:tblpPr')):
        tbl_pr.remove(pos)
    for wrap in tbl_pr.findall(qn('w:tblOverlap')):
        tbl_pr.remove(wrap)
    tbl_borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'nil')
        tbl_borders.append(b)
    tbl_pr.append(tbl_borders)
    for row in tbl.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            for old in tc_pr.findall(qn('w:tcBorders')):
                tc_pr.remove(old)
            tc_borders = OxmlElement('w:tcBorders')
            for edge in ('top', 'left', 'bottom', 'right'):
                b = OxmlElement(f'w:{edge}')
                b.set(qn('w:val'), 'nil')
                tc_borders.append(b)
            tc_pr.append(tc_borders)
            _clear_cell_shading(cell)


def _set_cell_margins(cell, top=40, right=60, bottom=40, left=60):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement('w:tcMar')
    for edge, val in (('top', top), ('left', left),
                      ('bottom', bottom), ('right', right)):
        m = OxmlElement(f'w:{edge}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tc_mar.append(m)
    tc_pr.append(tc_mar)


def _set_table_width(table, pct=100):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), f'{pct * 50}')
    tbl_w.set(qn('w:type'), 'pct')


def _set_table_layout_autofit(table):
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tbl_pr.append(layout)
    layout.set(qn('w:type'), 'autofit')


def _replace_status_emojis_in_paragraph(p):
    """Reemplaza 🟢🟡🔴 por ● coloreado en CADA run del párrafo.
    Si un run contiene un emoji, se divide: el texto antes/después
    queda en runs separados con el color heredado, y el ● va en su
    propio run con el color correspondiente."""
    for run in list(p.runs):
        text = run.text
        if not any(e in text for e in STATUS_EMOJI):
            continue
        # Reconstruir el texto carácter a carácter
        new_segments = []  # lista de (texto, color_or_None)
        buf = ''
        for ch in text:
            if ch in STATUS_EMOJI:
                if buf:
                    new_segments.append((buf, None))
                    buf = ''
                new_segments.append((STATUS_GLYPH, STATUS_EMOJI[ch]))
            else:
                buf += ch
        if buf:
            new_segments.append((buf, None))
        # Aplicar al run actual: primer segmento
        if new_segments:
            first_txt, first_color_hex = new_segments[0]
            run.text = first_txt
            if first_color_hex is not None:
                run.font.color.rgb = _hex_to_rgb(first_color_hex)
            # Insertar runs adicionales detrás del actual
            parent_p = run._element.getparent()
            insert_idx = list(parent_p).index(run._element)
            for seg_txt, seg_color_hex in new_segments[1:]:
                new_run = OxmlElement('w:r')
                orig_rpr = run._element.find(qn('w:rPr'))
                if orig_rpr is not None:
                    import copy as _copy
                    new_rpr = _copy.deepcopy(orig_rpr)
                    new_run.append(new_rpr)
                t = OxmlElement('w:t')
                t.text = seg_txt
                t.set(qn('xml:space'), 'preserve')
                new_run.append(t)
                insert_idx += 1
                parent_p.insert(insert_idx, new_run)
                if seg_color_hex is not None:
                    rpr = new_run.find(qn('w:rPr'))
                    if rpr is None:
                        rpr = OxmlElement('w:rPr')
                        new_run.insert(0, rpr)
                    for c in rpr.findall(qn('w:color')):
                        rpr.remove(c)
                    color_el = OxmlElement('w:color')
                    color_el.set(qn('w:val'), seg_color_hex)
                    rpr.append(color_el)


def _replace_status_emojis(doc):
    """Recorre body + celdas y aplica el reemplazo de emojis de estado."""
    for p in doc.paragraphs:
        _replace_status_emojis_in_paragraph(p)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_status_emojis_in_paragraph(p)


def _force_body_text_black(doc):
    """Recorre todo el documento (body + tablas) y fuerza color = negro,
    SALVO en runs cuyo color ya fue establecido en un valor de la paleta
    de estados (verde/amarillo/rojo) por el reemplazo de emojis."""
    status_hexes = {h.upper() for h in STATUS_EMOJI.values()}
    def _process_paragraph(p):
        for run in p.runs:
            rgb = run.font.color.rgb
            if rgb is not None and str(rgb).upper() in status_hexes:
                continue  # preservar color de estado
            run.font.color.rgb = BODY_TEXT
    for p in doc.paragraphs:
        _process_paragraph(p)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _process_paragraph(p)


def embellish(path: Path):
    print(f'──── {path.name}')
    if not path.exists():
        print('  SKIP (no existe)')
        return
    doc = Document(str(path))

    # 1) Reemplazar emojis 🟢🟡🔴 por ● coloreados ANTES de forzar negro
    _replace_status_emojis(doc)

    # 2) Forzar el resto del texto a negro (sin tocar los ● de estado)
    _force_body_text_black(doc)

    # 3) Desenvolver las tablas wrapper de figura (1x1 con imagen)
    unwrapped = _unwrap_figure_tables(doc)

    # 4) Estilizar cada tabla restante
    n = 0
    for tbl in doc.tables:
        n += 1
        _strip_table_style(tbl)
        _set_table_width(tbl, pct=100)
        _set_table_layout_autofit(tbl)
        for r_idx, row in enumerate(tbl.rows):
            for c_idx, cell in enumerate(row.cells):
                _set_cell_borders(cell)
                _set_cell_margins(cell)
                _clear_cell_shading(cell)
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after  = Pt(0)
                    for run in p.runs:
                        run.font.size = TABLE_FONT_SZ
                        if r_idx == 0:
                            # Header: solo bold, sin fondo, texto negro
                            run.font.bold = True
                            run.font.color.rgb = BODY_TEXT

    # 5) Quitar bordes de párrafos (captions) y figuras
    _strip_paragraph_borders(doc)
    _strip_caption_shading(doc)
    _strip_figure_borders(doc)
    # 6) Compactar layout (margenes + fuente Normal) para que el DOCX
    #    se pagine cerca del PDF en vez de duplicar paginas
    _compact_layout(doc)

    doc.save(str(path))
    print(f'  ✓ {n} tablas embellecidas · {unwrapped} wrappers de figura eliminados')


if __name__ == '__main__':
    if len(sys.argv) < 2:
        ROOT = Path(__file__).resolve().parents[1]
        targets = [
            ROOT / 'docs' / 'informe_final.docx',
            ROOT / 'docs' / 'informe_anexos.docx',
        ]
    else:
        targets = [Path(a) for a in sys.argv[1:]]
    for t in targets:
        embellish(t)
    print('\nListo. Abrir los DOCX para ver el resultado.')
